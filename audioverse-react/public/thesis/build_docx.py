#!/usr/bin/env python3
"""
build_docx.py — Łączy pliki rozdziałów 1–5 (.txt, format Markdown-like)
w jeden dokument .docx sformatowany pod wytyczne UBB.

Obsługuje:
  - nagłówki # / ## / ###  →  Heading 1 / 2 / 3
  - **bold**, *italic*, ***bold-italic***
  - odnośniki [N]  →  przypisy dolne (footnotes)
  - tabele Markdown  →  tabele Word z obramowaniem
  - bloki kodu ```lang ... ```  →  szare tło, czcionka Consolas 8pt
  - blockquotes > ...  →  kursywa z wcięciem (opisy rysunków/tabel)
  - listy numerowane (1. / 2.) i punktowane (- / *)
  - równania LaTeX $...$ i $$...$$  →  kursywa Cambria Math (uproszczenie)

Zależności:  pip install python-docx
Uruchomienie:
  cd documentation/Seminarium dyplomowe
  python build_docx.py
"""
from __future__ import annotations

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import latex2mathml.converter
import lxml.etree as etree

# ──────────────────────────────────────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
INPUT_FILES = [
    SCRIPT_DIR / f"Gra wokalna karaoke z wykorzystaniem sztucznej inteligencji - {i}.txt"
    for i in range(1, 6)
]
OUTPUT_FILE = SCRIPT_DIR / "Gra wokalna karaoke z wykorzystaniem sztucznej inteligencji.docx"

FONT_MAIN = "Times New Roman"
FONT_CODE = "Consolas"
FONT_MATH = "Cambria Math"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_CODE = Pt(8)
LINE_SPACING = 1.5

NBSP = "\u00A0"  # non-breaking space (twarda spacja)
GRAPHICS_DIR = SCRIPT_DIR / "graphics"

# ──────────────────────────────────────────────────────────────────────────────
# Bibliografia — mapa [N] → pełny tekst przypisu
# ──────────────────────────────────────────────────────────────────────────────

_bibliography: dict[str, str] = {}   # populated by _load_bibliography()

# Collectors for Spis ilustracji / Spis tabel
_collected_figures: list[tuple[str, str]] = []  # (num, caption) e.g. ("1.1", "Główny ekran...")
_collected_tables: list[tuple[str, str]] = []   # (num, caption) e.g. ("4.1", "Wymagania...")

_first_h1_seen: bool = False  # page break only before the very first chapter heading


def _load_bibliography():
    """Scan all input .txt files for a bibliography section and parse [N] entries.

    Entries may span multiple lines (continuation lines start with whitespace).
    We strip markdown italics (*...*) and „..." quotes so the footnote text is
    plain.  Result stored in module-level _bibliography dict.
    """
    global _bibliography
    bib: dict[str, str] = {}
    for fp in INPUT_FILES:
        if not fp.exists():
            continue
        lines = fp.read_text(encoding="utf-8").split("\n")
        in_bib = False
        current_num: str | None = None
        current_parts: list[str] = []

        for line in lines:
            stripped = line.strip()

            # Detect start of bibliography section
            if stripped in ("## Literatura", "BIBLIOGRAFIA", "## Bibliografia"):
                in_bib = True
                continue

            if not in_bib:
                continue

            # Detect end — next heading or thematic break
            if stripped.startswith("#") or re.match(r"^-{3,}$", stripped):
                # save pending
                if current_num and current_parts:
                    bib[current_num] = _clean_bib_entry(" ".join(current_parts))
                in_bib = False
                current_num = None
                current_parts = []
                continue

            # New entry: [N]
            m = re.match(r"^\[(\d{1,3})\]\s+(.+)", stripped)
            if m:
                # save previous
                if current_num and current_parts:
                    bib[current_num] = _clean_bib_entry(" ".join(current_parts))
                current_num = m.group(1)
                current_parts = [m.group(2)]
                continue

            # Continuation line (indented)
            if current_num and stripped:
                current_parts.append(stripped)

        # End of file — save last entry
        if current_num and current_parts:
            bib[current_num] = _clean_bib_entry(" ".join(current_parts))

    _bibliography = bib


def _clean_bib_entry(text: str) -> str:
    """Strip markdown italics and normalise whitespace."""
    text = text.replace("*", "")      # remove markdown italics
    text = text.replace("\u201e", "\"").replace("\u201d", "\"")  # „ " → " "
    text = re.sub(r"\s+", " ", text).strip()
    if not text.endswith("."):
        text += "."
    return text


def get_citation_text(num: str) -> str:
    """Return full bibliography text for reference [num], or a fallback."""
    return _bibliography.get(num, f"Patrz: pozycja [{num}] w bibliografii.")


def _classify_bib_entry(text: str) -> str:
    """Classify a bibliography entry as 'book', 'article', or 'web'.

    Heuristic:
    - If text contains „..." (article title in Polish „...") → article
    - If text contains http:// or https:// → web
    - Otherwise → book
    """
    if "\u201e" in text or '\"' in text:
        # Has quoted title → article or conference paper
        return "article"
    if "http://" in text or "https://" in text:
        return "web"
    return "book"


def _build_bibliography(doc: Document):
    """Build a properly formatted bibliography section grouped by type."""
    if not _bibliography:
        return

    print("\n  Budowanie sekcji bibliografii...")

    _add_unnumbered_heading(doc, "Bibliografia")

    # Classify entries
    books: list[tuple[int, str]] = []
    articles: list[tuple[int, str]] = []
    web: list[tuple[int, str]] = []

    for num_str, text in _bibliography.items():
        num = int(num_str)
        cat = _classify_bib_entry(text)
        if cat == "book":
            books.append((num, text))
        elif cat == "article":
            articles.append((num, text))
        else:
            web.append((num, text))

    # Sort each group by original number
    books.sort(key=lambda x: x[0])
    articles.sort(key=lambda x: x[0])
    web.sort(key=lambda x: x[0])

    seq = 1  # sequential numbering from 1

    def _add_group(group_title: str, items: list[tuple[int, str]]):
        nonlocal seq
        if not items:
            return
        _add_unnumbered_heading(doc, group_title, level=2, page_break=False)
        for _orig_num, entry_text in items:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)
            run_num = p.add_run(f"[{seq}]  ")
            run_num.bold = True
            run_num.font.name = FONT_MAIN
            run_num.font.size = FONT_SIZE_BODY
            run_text = p.add_run(entry_text)
            run_text.font.name = FONT_MAIN
            run_text.font.size = FONT_SIZE_BODY
            seq += 1

    _add_group("Książki i monografie", books)
    _add_group("Artykuły naukowe i konferencyjne", articles)
    _add_group("Źródła internetowe", web)

    print(f"  ✓ Bibliografia: {len(books)} książek, {len(articles)} artykułów, {len(web)} źródeł internetowych")


def _add_unnumbered_heading(doc: Document, text: str, *, level: int = 1,
                            page_break: bool = True):
    """Add a heading-like paragraph WITHOUT Word auto-numbering.

    Used for structural sections like "Spis treści", "Spis ilustracji",
    "Spis tabel" that should not carry chapter numbers.
    """
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24) if level == 1 else Pt(18)
    p.paragraph_format.space_after = Pt(12) if level == 1 else Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_MAIN
    run.font.size = Pt(16) if level == 1 else Pt(14)
    return p


def _build_list_of_figures(doc: Document):
    """Build 'Spis ilustracji' from collected figure captions."""
    if not _collected_figures:
        return
    print("\n  Budowanie spisu ilustracji...")
    _add_unnumbered_heading(doc, "Spis ilustracji")

    for fig_num, caption in _collected_figures:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(-1.5)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        # Clean caption — remove "Rysunek N.M — " prefix if already in caption
        clean = re.sub(r"^Rysunek\s+\d+\.\d+\s*[—–\-]\s*", "", caption)
        run = p.add_run(f"Rysunek {fig_num}. ")
        run.bold = True
        run.font.name = FONT_MAIN
        run.font.size = FONT_SIZE_BODY
        run2 = p.add_run(clean)
        run2.font.name = FONT_MAIN
        run2.font.size = FONT_SIZE_BODY

    print(f"  ✓ Spis ilustracji: {len(_collected_figures)} pozycji")


def _build_list_of_tables(doc: Document):
    """Build 'Spis tabel' from collected table captions."""
    if not _collected_tables:
        return
    print("\n  Budowanie spisu tabel...")
    _add_unnumbered_heading(doc, "Spis tabel")

    for tab_num, caption in _collected_tables:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(-1.5)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        # Clean caption — remove "Tabela N.M. " prefix if already in caption
        clean = re.sub(r"^Tabela\s+\d+\.\d+\.?\s*", "", caption)
        run = p.add_run(f"Tabela {tab_num}. ")
        run.bold = True
        run.font.name = FONT_MAIN
        run.font.size = FONT_SIZE_BODY
        run2 = p.add_run(clean)
        run2.font.name = FONT_MAIN
        run2.font.size = FONT_SIZE_BODY

    print(f"  ✓ Spis tabel: {len(_collected_tables)} pozycji")


# ──────────────────────────────────────────────────────────────────────────────
# LaTeX → OMML (Office Math) conversion
# ──────────────────────────────────────────────────────────────────────────────

# Load the MML2OMML.XSL stylesheet shipped with Microsoft Office
_MML2OMML_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
_omml_xslt: etree.XSLT | None = None


def _get_omml_xslt() -> etree.XSLT | None:
    """Lazy-load the XSLT transform (once)."""
    global _omml_xslt
    if _omml_xslt is not None:
        return _omml_xslt
    if not Path(_MML2OMML_PATH).exists():
        print(f"  ⚠ MML2OMML.XSL not found at {_MML2OMML_PATH} — equations will be plain text")
        return None
    _omml_xslt = etree.XSLT(etree.parse(_MML2OMML_PATH))
    return _omml_xslt


def _latex_to_omml(latex: str) -> etree._Element | None:
    """Convert a LaTeX string to an <m:oMath> OMML element, or None on failure."""
    xslt = _get_omml_xslt()
    if xslt is None:
        return None
    try:
        mathml_str = latex2mathml.converter.convert(latex)
        mml_tree = etree.fromstring(mathml_str.encode("utf-8"))
        omml_tree = xslt(mml_tree)
        return omml_tree.getroot()
    except Exception as exc:
        print(f"  ⚠ Equation conversion failed for '{latex[:40]}…': {exc}")
        return None


def add_omml_to_paragraph(paragraph, latex: str):
    """Insert an OMML equation into an existing paragraph. Falls back to italic Cambria Math."""
    omml = _latex_to_omml(latex)
    if omml is not None:
        paragraph._p.append(omml)
    else:
        # Fallback: plain italic text
        run = paragraph.add_run(latex)
        run.italic = True
        run.font.name = FONT_MATH
        run.font.size = FONT_SIZE_BODY


def add_display_equation(doc: Document, latex: str):
    """Add a centred display equation as an <m:oMathPara> block."""
    omml = _latex_to_omml(latex)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    if omml is not None:
        # Wrap in <m:oMathPara> for proper centred display
        MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        oMathPara = etree.SubElement(p._p, f"{{{MATH_NS}}}oMathPara")
        oMathPara.append(omml)
    else:
        run = p.add_run(latex)
        run.italic = True
        run.font.name = FONT_MATH
        run.font.size = FONT_SIZE_BODY

# ──────────────────────────────────────────────────────────────────────────────
# Mapa: numer rysunku → numer pliku graficznego (001–030)
# ──────────────────────────────────────────────────────────────────────────────
FIGURE_FILE_MAP: dict[str, str] = {
    # rys. 1.1 usunięty (screeny z implementacji nie należą do rozdziału 1)
    "2.1": "002", "2.2": "017", "2.3": "018", "2.4": "019", "2.5": "003",
    "2.6": "020",  # Spektrogram STFT i Mel
    "2.7": "004",  # Porównanie trajektorii F0 pYIN vs CREPE
    # Rozdział 3: 3.1 = schemat komunikacji (pierwszy w tekście), 3.2 = diagram ER
    "3.1": "005",  # Schemat komunikacji między warstwami
    "3.2": "007",  # Diagram encji (ER)
    # Rozdział 4: CQRS jako 4.1 (nie ma osobnego rys. architektury w tekście)
    "4.1": "021",  # CQRS z MediatR
    "4.2": "022",  # Przepływ sesji karaoke
    "4.3": "023",  # Autoryzacja JWT i OAuth 2.0
    "4.4": "008",  # Swagger UI modułu laboratoryjnego
    "4.5": "024",  # Pipeline modułu laboratoryjnego
    "4.6": "025",  # Struktura routingu frontend
    "4.7": "009",  # Przeglądarka piosenek
    "4.8": "010",  # Ekran rozgrywki karaoke
    "4.9": "011",  # Ekran rankingu
    "4.10": "026", # Tworzenie sesji / dołączanie
    "4.11": "027", # Ustawienia wejścia audio
    "4.12": "028", # Strona modułu laboratoryjnego (admin)
    "4.13": "012", # Diagram kontenerów Docker Compose
    "5.1": "029", "5.2": "013", "5.3": "030",  # 5.4 i 5.5 usunięte wraz z sekcjami
}


def _read_figure_source(figure_num: str) -> tuple[str, str, str] | None:
    """Read NNN_source.txt for the given figure number.

    Returns (site_name, url, access_date) or None if no source file exists.
    Format of NNN_source.txt:
      Line 1: site name
      Line 2: URL
      Line 3: access date (YYYY-MM-DD)
    """
    file_num = FIGURE_FILE_MAP.get(figure_num)
    if not file_num:
        return None
    source_path = GRAPHICS_DIR / f"{file_num}_source.txt"
    if not source_path.exists():
        return None
    lines = source_path.read_text(encoding="utf-8").strip().splitlines()
    if len(lines) < 3:
        return None
    return lines[0].strip(), lines[1].strip(), lines[2].strip()


def _find_figure_image(figure_num: str) -> Path | None:
    """Find the image file (png or jpg) for a given figure number."""
    file_num = FIGURE_FILE_MAP.get(figure_num)
    if not file_num:
        return None
    for ext in ("png", "jpg", "jpeg"):
        img_path = GRAPHICS_DIR / f"{file_num}.{ext}"
        if img_path.exists():
            return img_path
    return None


def _insert_figure_image(doc: Document, figure_num: str):
    """Insert the image for a figure if it exists. Returns True if inserted."""
    img_path = _find_figure_image(figure_num)
    if img_path is None:
        # No image file — insert placeholder text
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"[Rysunek {figure_num} — brak pliku graficznego]")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        return False

    # Determine image dimensions to handle portrait images gracefully.
    # Max width  = 14 cm (text area); max height = ~13 cm (~half A4 page).
    # For portrait images that would exceed 13 cm tall at full width,
    # constrain by height instead so they don't dominate the page.
    MAX_W_CM = 14.0
    MAX_H_CM = 13.0   # ~half of 24.7 cm text area

    img_width_arg = Cm(MAX_W_CM)
    img_height_arg = None

    try:
        from PIL import Image as _PILImage
        with _PILImage.open(str(img_path)) as _im:
            px_w, px_h = _im.size
        if px_w > 0 and px_h > 0:
            aspect = px_h / px_w          # > 1 means portrait
            natural_h_cm = aspect * MAX_W_CM
            if natural_h_cm > MAX_H_CM:
                # Constrain by height — do not set width so Word scales correctly
                img_width_arg = None
                img_height_arg = Cm(MAX_H_CM)
    except Exception:
        pass  # fall back to fixed width if PIL unavailable or image unreadable

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(img_path), width=img_width_arg, height=img_height_arg)
    print(f"    ✓ Wstawiono obraz: {img_path.name} (Rysunek {figure_num})")
    return True

# ──────────────────────────────────────────────────────────────────────────────
# Zasada sierot — twarde spacje po krótkich wyrazach
# ──────────────────────────────────────────────────────────────────────────────
# W polskiej typografii akademickiej jednoliterowe spójniki / przyimki
# (i, a, o, u, w, z, e) oraz krótkie dwuliterowe (do, na, od, po, we, za,
# ze, że, by, to, co, bo, ku, ni, no, aż, al) nie mogą zostawać samotnie
# na końcu wiersza.  Rozwiązanie: po takim wyrazie stawiamy twardą spację
# (U+00A0 non-breaking space), która zabrania łamania linii w tym miejscu.
# ──────────────────────────────────────────────────────────────────────────────

_NBSP_WORDS_LOWER = {
    # jednoliterowe — zawsze
    "a", "i", "o", "u", "w", "z", "e",
    # dwuliterowe przyimki / spójniki / partykuły
    "do", "na", "od", "po", "we", "za", "ze",
    "ni", "ku", "no", "że", "by", "to", "co",
    "bo", "al", "aż",
}

# Regex: granica słowa + wyraz ze słownika + zwykła spacja
# Budujemy wariant case-insensitive, więc wystarczą formy małe
_RE_NBSP = re.compile(
    r"\b(" + "|".join(sorted(_NBSP_WORDS_LOWER, key=len, reverse=True)) + r")\s",
    re.IGNORECASE,
)


def apply_nbsp(text: str) -> str:
    """Zamień spacje po krótkich wyrazach na twarde spacje (zasada sierot)."""
    return _RE_NBSP.sub(lambda m: m.group(1) + NBSP, text)

# ──────────────────────────────────────────────────────────────────────────────
# List numbering setup (bullet / numbered lists)
# ──────────────────────────────────────────────────────────────────────────────

def _wire_list_numbering(doc: Document, *, bullet: bool = True, number: bool = True):
    """Create proper numbering definitions for List Bullet / List Number styles.

    python-docx doesn't expose a high-level API for numbering, so we
    manipulate the numbering XML part directly.

    IMPORTANT: When the numbering part is an XmlPart (loaded from a .docx),
    its blob property REGENERATES content from _element — setting _blob is a
    no-op.  We must work directly on _element instead of creating a copy via
    etree.fromstring(part.blob).
    """
    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Find or create numbering part
    numbering_part = None
    for rel in doc.part.rels.values():
        if "numbering" in rel.reltype:
            numbering_part = rel.target_part
            break

    create_new_part = numbering_part is None
    if not create_new_part:
        # Work DIRECTLY on the part's live element so changes are reflected on save.
        # Do NOT use etree.fromstring(part.blob) — that creates a detached copy.
        numbering_el = numbering_part._element
    else:
        numbering_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '</w:numbering>'
        )
        numbering_el = etree.fromstring(numbering_xml.encode("utf-8"))

    # Determine next available IDs
    existing_abstract = numbering_el.findall(f"{{{W}}}abstractNum")
    existing_num = numbering_el.findall(f"{{{W}}}num")
    next_abstract_id = max((int(a.get(f"{{{W}}}abstractNumId", "0")) for a in existing_abstract), default=-1) + 1
    next_num_id = max((int(n.get(f"{{{W}}}numId", "0")) for n in existing_num), default=0) + 1

    # Find insertion point: abstractNum elements must precede num elements in
    # the OOXML schema.  Insert new abstractNum nodes just before the first
    # existing <w:num> so ordering stays valid.
    all_children = list(numbering_el)
    abstract_insert_pos = len(all_children)
    for i, child in enumerate(all_children):
        if child.tag == f"{{{W}}}num":
            abstract_insert_pos = i
            break

    bullet_num_id = None
    number_num_id = None

    def _build_abstract_bullet(abstract_id: int):
        abstract = etree.Element(f"{{{W}}}abstractNum")
        abstract.set(f"{{{W}}}abstractNumId", str(abstract_id))
        lvl = etree.SubElement(abstract, f"{{{W}}}lvl")
        lvl.set(f"{{{W}}}ilvl", "0")
        start = etree.SubElement(lvl, f"{{{W}}}start")
        start.set(f"{{{W}}}val", "1")
        numFmt = etree.SubElement(lvl, f"{{{W}}}numFmt")
        numFmt.set(f"{{{W}}}val", "bullet")
        lvlText = etree.SubElement(lvl, f"{{{W}}}lvlText")
        lvlText.set(f"{{{W}}}val", "\uf0b7")  # Symbol private-use bullet (U+F0B7)
        lvlJc = etree.SubElement(lvl, f"{{{W}}}lvlJc")
        lvlJc.set(f"{{{W}}}val", "left")
        pPr = etree.SubElement(lvl, f"{{{W}}}pPr")
        ind = etree.SubElement(pPr, f"{{{W}}}ind")
        ind.set(f"{{{W}}}left", "720")    # 720 twips = 1.27 cm
        ind.set(f"{{{W}}}hanging", "360") # 360 twips = 0.63 cm
        rPr = etree.SubElement(lvl, f"{{{W}}}rPr")
        rFonts = etree.SubElement(rPr, f"{{{W}}}rFonts")
        rFonts.set(f"{{{W}}}ascii", "Symbol")
        rFonts.set(f"{{{W}}}hAnsi", "Symbol")
        rFonts.set(f"{{{W}}}hint", "default")
        return abstract

    def _build_abstract_number(abstract_id: int):
        abstract = etree.Element(f"{{{W}}}abstractNum")
        abstract.set(f"{{{W}}}abstractNumId", str(abstract_id))
        lvl = etree.SubElement(abstract, f"{{{W}}}lvl")
        lvl.set(f"{{{W}}}ilvl", "0")
        start = etree.SubElement(lvl, f"{{{W}}}start")
        start.set(f"{{{W}}}val", "1")
        numFmt = etree.SubElement(lvl, f"{{{W}}}numFmt")
        numFmt.set(f"{{{W}}}val", "decimal")
        lvlText = etree.SubElement(lvl, f"{{{W}}}lvlText")
        lvlText.set(f"{{{W}}}val", "%1.")
        lvlJc = etree.SubElement(lvl, f"{{{W}}}lvlJc")
        lvlJc.set(f"{{{W}}}val", "left")
        pPr = etree.SubElement(lvl, f"{{{W}}}pPr")
        ind = etree.SubElement(pPr, f"{{{W}}}ind")
        ind.set(f"{{{W}}}left", "720")
        ind.set(f"{{{W}}}hanging", "360")
        return abstract

    if bullet:
        abstract = _build_abstract_bullet(next_abstract_id)
        numbering_el.insert(abstract_insert_pos, abstract)
        abstract_insert_pos += 1  # next abstractNum goes after this one

        num_el = etree.SubElement(numbering_el, f"{{{W}}}num")
        num_el.set(f"{{{W}}}numId", str(next_num_id))
        abstractNumIdRef = etree.SubElement(num_el, f"{{{W}}}abstractNumId")
        abstractNumIdRef.set(f"{{{W}}}val", str(next_abstract_id))
        bullet_num_id = next_num_id

        next_abstract_id += 1
        next_num_id += 1

    if number:
        abstract = _build_abstract_number(next_abstract_id)
        numbering_el.insert(abstract_insert_pos, abstract)
        abstract_insert_pos += 1

        num_el = etree.SubElement(numbering_el, f"{{{W}}}num")
        num_el.set(f"{{{W}}}numId", str(next_num_id))
        abstractNumIdRef = etree.SubElement(num_el, f"{{{W}}}abstractNumId")
        abstractNumIdRef.set(f"{{{W}}}val", str(next_abstract_id))
        number_num_id = next_num_id

        next_abstract_id += 1
        next_num_id += 1

    # For a newly created numbering part (no existing part), persist it now.
    # For an existing XmlPart we modified _element in-place — no write-back needed.
    if create_new_part:
        numbering_blob = etree.tostring(numbering_el, xml_declaration=True,
                                         encoding="UTF-8", standalone=True)
        numbering_part = OpcPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_blob,
            doc.part.package,
        )
        doc.part.relate_to(
            numbering_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
        )

    # Wire the styles to their numbering definitions via numPr
    if bullet and bullet_num_id is not None:
        _set_style_numPr(doc.styles["List Bullet"], bullet_num_id, ilvl=0)
    if number and number_num_id is not None:
        _set_style_numPr(doc.styles["List Number"], number_num_id, ilvl=0)


def _set_style_numPr(style, num_id: int, ilvl: int = 0):
    """Set <w:numPr> on a style's paragraph properties."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    pPr = style.element.find(f"{{{W}}}pPr")
    if pPr is None:
        pPr = etree.SubElement(style.element, f"{{{W}}}pPr")
    # Remove old numPr if present
    old = pPr.find(f"{{{W}}}numPr")
    if old is not None:
        pPr.remove(old)
    numPr = etree.SubElement(pPr, f"{{{W}}}numPr")
    ilvl_el = etree.SubElement(numPr, f"{{{W}}}ilvl")
    ilvl_el.set(f"{{{W}}}val", str(ilvl))
    numId_el = etree.SubElement(numPr, f"{{{W}}}numId")
    numId_el.set(f"{{{W}}}val", str(num_id))


# ──────────────────────────────────────────────────────────────────────────────
# Document setup
# ──────────────────────────────────────────────────────────────────────────────

def create_document(base_docx: Path | None = None) -> Document:
    """Create a document with A4 margins and base styles.
    
    If *base_docx* is given, opens that file as the starting document
    (preserving its content — e.g. a title page designed in Word).
    Otherwise creates a blank document.
    """
    if base_docx and base_docx.exists():
        doc = Document(str(base_docx))
        print(f"  ✓ Załadowano bazowy dokument: {base_docx.name}")
    else:
        doc = Document()

    # Page setup — A4, margins 2.5 cm
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)   # left margin wider for binding
    section.right_margin = Cm(2.0)

    # Default paragraph style — only override when building from scratch.
    # When loading TitlePage.docx we intentionally leave the Normal style
    # untouched so the title page layout stays intact.
    if not (base_docx and base_docx.exists()):
        style = doc.styles["Normal"]
        font = style.font
        font.name = FONT_MAIN
        font.size = FONT_SIZE_BODY
        pf = style.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_after = Pt(6)
        pf.first_line_indent = Cm(1.25)

    # Heading styles — numbered chapters
    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hfont = hstyle.font
        hfont.name = FONT_MAIN
        hfont.bold = True
        hfont.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            hfont.size = Pt(16)
            hstyle.paragraph_format.space_before = Pt(24)
            hstyle.paragraph_format.space_after = Pt(12)
            # page_break_before is applied directly on each chapter heading
            # paragraph (in process_file) to avoid breaking the title page.
        elif level == 2:
            hfont.size = Pt(14)
            hstyle.paragraph_format.space_before = Pt(18)
            hstyle.paragraph_format.space_after = Pt(8)
        else:
            hfont.size = Pt(12)
            hstyle.paragraph_format.space_before = Pt(12)
            hstyle.paragraph_format.space_after = Pt(6)
        hstyle.paragraph_format.first_line_indent = Cm(0)

        # Remove Word auto-numbering from heading styles
        # (TitlePage.docx may define multi-level list numbering which causes
        #  double numbering since our markdown headings already contain numbers)
        style_pPr = hstyle.element.find(f"{{{_W_NS}}}pPr")
        if style_pPr is not None:
            numPr = style_pPr.find(f"{{{_W_NS}}}numPr")
            if numPr is not None:
                style_pPr.remove(numPr)

    # Footnote text style — 9pt, single spacing, no indent
    try:
        fn_style = doc.styles["Footnote Text"]
    except KeyError:
        fn_style = doc.styles.add_style("Footnote Text", 1)  # 1 = paragraph
    fn_style.font.name = FONT_MAIN
    fn_style.font.size = Pt(9)
    fn_style.paragraph_format.line_spacing = 1.0
    fn_style.paragraph_format.space_after = Pt(2)
    fn_style.paragraph_format.space_before = Pt(0)
    fn_style.paragraph_format.first_line_indent = Cm(0)

    # Footnote reference character style — superscript
    try:
        fnr_style = doc.styles["Footnote Reference"]
    except KeyError:
        fnr_style = doc.styles.add_style("Footnote Reference", 2)  # 2 = character
    fnr_style.font.superscript = True
    fnr_style.font.size = Pt(9)

    # List styles — ensure they exist (may be missing if loading from a custom .docx)
    # When creating from scratch, we must also wire up numbering definitions
    # so that Word actually renders bullet / number characters.
    from docx.enum.style import WD_STYLE_TYPE

    try:
        _ = doc.styles["List Bullet"]
    except KeyError:
        ls = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        ls.base_style = doc.styles["Normal"]
        ls.font.name = FONT_MAIN
        ls.font.size = FONT_SIZE_BODY
        ls.paragraph_format.left_indent = Cm(1.25)
        ls.paragraph_format.first_line_indent = Cm(0)

    try:
        _ = doc.styles["List Number"]
    except KeyError:
        ls = doc.styles.add_style("List Number", WD_STYLE_TYPE.PARAGRAPH)
        ls.base_style = doc.styles["Normal"]
        ls.font.name = FONT_MAIN
        ls.font.size = FONT_SIZE_BODY
        ls.paragraph_format.left_indent = Cm(1.25)
        ls.paragraph_format.first_line_indent = Cm(0)

    # Always re-wire numbering definitions so the template's own numbering
    # (which may use decimal format for both styles) is overridden correctly.
    _wire_list_numbering(doc, bullet=True, number=True)

    return doc

_footnote_counter = 0

def _ensure_footnotes_part(doc: Document):
    """Create the footnotes XML part if it doesn't exist yet."""
    package = doc.part.package
    # Check if footnotes part already registered
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            return
    # Create footnotes part
    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI
    import lxml.etree as etree

    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        '  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '  <w:footnote w:type="separator" w:id="-1">'
        '    <w:p><w:r><w:separator/></w:r></w:p>'
        '  </w:footnote>'
        '  <w:footnote w:type="continuationSeparator" w:id="0">'
        '    <w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '  </w:footnote>'
        '</w:footnotes>'
    )
    footnotes_part = OpcPart(
        PackURI("/word/footnotes.xml"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        footnotes_xml.encode("utf-8"),
        package,
    )
    doc.part.relate_to(
        footnotes_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
    )


def _get_footnotes_element(doc: Document):
    """Get the <w:footnotes> root element."""
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            import lxml.etree as etree
            return etree.fromstring(rel.target_part.blob)
    return None


def _save_footnotes(doc: Document, footnotes_el):
    """Persist modified footnotes XML back to the part."""
    import lxml.etree as etree
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            rel.target_part._blob = etree.tostring(footnotes_el, xml_declaration=True, encoding="UTF-8", standalone=True)
            return


def add_footnote_ref(paragraph, doc: Document, footnote_text: str):
    """Insert a footnote reference in the paragraph and add footnote content.

    The reference number is rendered as superscript in the body text.
    The footnote itself appears in the footnote area at the bottom of the page.
    """
    global _footnote_counter
    _ensure_footnotes_part(doc)

    _footnote_counter += 1
    fn_id = _footnote_counter

    # Add footnote to the footnotes part
    footnotes_el = _get_footnotes_element(doc)
    if footnotes_el is None:
        # Fallback: just add bracketed text inline
        run = paragraph.add_run(f" [{footnote_text}]")
        run.font.size = Pt(9)
        return

    fn = OxmlElement("w:footnote")
    fn.set(qn("w:id"), str(fn_id))

    fn_p = OxmlElement("w:p")
    fn_pPr = OxmlElement("w:pPr")
    fn_pStyle = OxmlElement("w:pStyle")
    fn_pStyle.set(qn("w:val"), "FootnoteText")
    fn_pPr.append(fn_pStyle)
    # Single line spacing, 9pt
    fn_spacing = OxmlElement("w:spacing")
    fn_spacing.set(qn("w:after"), "40")     # ~2pt
    fn_spacing.set(qn("w:line"), "240")     # single spacing
    fn_spacing.set(qn("w:lineRule"), "auto")
    fn_pPr.append(fn_spacing)
    fn_p.append(fn_pPr)

    # Footnote auto-number marker inside the footnote body (superscript)
    fn_r_ref = OxmlElement("w:r")
    fn_rPr_ref = OxmlElement("w:rPr")
    fn_rStyle_ref = OxmlElement("w:rStyle")
    fn_rStyle_ref.set(qn("w:val"), "FootnoteReference")
    fn_rPr_ref.append(fn_rStyle_ref)
    # Explicit superscript + size for the number in footnote body
    fn_vertAlign = OxmlElement("w:vertAlign")
    fn_vertAlign.set(qn("w:val"), "superscript")
    fn_rPr_ref.append(fn_vertAlign)
    fn_sz_ref = OxmlElement("w:sz")
    fn_sz_ref.set(qn("w:val"), "18")  # 9pt
    fn_rPr_ref.append(fn_sz_ref)
    fn_r_ref.append(fn_rPr_ref)
    fn_ref_mark = OxmlElement("w:footnoteRef")
    fn_r_ref.append(fn_ref_mark)
    fn_p.append(fn_r_ref)

    # Space + footnote body text (9pt, Times New Roman)
    fn_r_text = OxmlElement("w:r")
    fn_rPr_text = OxmlElement("w:rPr")
    fn_sz = OxmlElement("w:sz")
    fn_sz.set(qn("w:val"), "18")  # 9pt
    fn_rPr_text.append(fn_sz)
    fn_szCs = OxmlElement("w:szCs")
    fn_szCs.set(qn("w:val"), "18")
    fn_rPr_text.append(fn_szCs)
    fn_rFont = OxmlElement("w:rFonts")
    fn_rFont.set(qn("w:ascii"), FONT_MAIN)
    fn_rFont.set(qn("w:hAnsi"), FONT_MAIN)
    fn_rPr_text.append(fn_rFont)
    fn_r_text.append(fn_rPr_text)
    fn_t = OxmlElement("w:t")
    fn_t.set(qn("xml:space"), "preserve")
    fn_t.text = f" {footnote_text}"
    fn_r_text.append(fn_t)
    fn_p.append(fn_r_text)

    fn.append(fn_p)
    footnotes_el.append(fn)
    _save_footnotes(doc, footnotes_el)

    # ── Insert superscript footnote reference in the main body text ──
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "FootnoteReference")
    rPr.append(rStyle)
    # Explicit superscript so Word renders it correctly even without
    # the built-in FootnoteReference style definition
    vertAlign = OxmlElement("w:vertAlign")
    vertAlign.set(qn("w:val"), "superscript")
    rPr.append(vertAlign)
    r.append(rPr)
    fn_ref = OxmlElement("w:footnoteReference")
    fn_ref.set(qn("w:id"), str(fn_id))
    r.append(fn_ref)
    paragraph._p.append(r)


# ──────────────────────────────────────────────────────────────────────────────
# Inline formatting helpers
# ──────────────────────────────────────────────────────────────────────────────

# Regex for inline elements (order matters — bold-italic first)
_RE_INLINE = re.compile(
    r"(\*\*\*(.+?)\*\*\*)"           # group 1,2: bold-italic
    r"|(\*\*(.+?)\*\*)"              # group 3,4: bold
    r"|(\*(.+?)\*)"                  # group 5,6: italic
    r"|(`([^`]+)`)"                  # group 7,8: inline code
    r"|(\$\$(.+?)\$\$)"             # group 9,10: display math
    r"|(\$([^$]+?)\$)"              # group 11,12: inline math
    r"|(\s?\[(\d{1,3})\])"          # group 13,14: citation [N] (eat optional preceding space)
    r"|(\{EN:\s*(.+?)\})"           # group 15,16: {EN: original text} → footnote with original
)


def add_formatted_text(paragraph, text: str, doc: Document, *, base_bold=False, base_italic=False):
    """Parse inline markdown in *text* and add runs to *paragraph*."""
    text = apply_nbsp(text)  # zasada sierot — twarde spacje
    pos = 0
    for m in _RE_INLINE.finditer(text):
        # plain text before this match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = FONT_MAIN
            run.font.size = FONT_SIZE_BODY
            if base_bold:
                run.bold = True
            if base_italic:
                run.italic = True

        if m.group(2):       # ***bold-italic***
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
            run.font.name = FONT_MAIN
            run.font.size = FONT_SIZE_BODY
        elif m.group(4):     # **bold**
            run = paragraph.add_run(m.group(4))
            run.bold = True
            run.font.name = FONT_MAIN
            run.font.size = FONT_SIZE_BODY
        elif m.group(6):     # *italic*
            run = paragraph.add_run(m.group(6))
            run.italic = True
            run.font.name = FONT_MAIN
            run.font.size = FONT_SIZE_BODY
        elif m.group(8):     # `inline code`
            run = paragraph.add_run(m.group(8))
            run.font.name = FONT_CODE
            run.font.size = FONT_SIZE_CODE
            # light gray background via shading
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "E8E8E8")
            run._r.get_or_add_rPr().append(shd)
        elif m.group(10):    # $$display math$$ (inline occurrence)
            add_omml_to_paragraph(paragraph, m.group(10))
        elif m.group(12):    # $inline math$
            add_omml_to_paragraph(paragraph, m.group(12))
        elif m.group(14):    # [N] citation → footnote
            citation_num = m.group(14)
            add_footnote_ref(paragraph, doc, get_citation_text(citation_num))
        elif m.group(16):    # {EN: original text} → footnote with original English
            en_text = m.group(16).strip()
            add_footnote_ref(paragraph, doc, f"Tekst oryginału: \"{en_text}\".")

        pos = m.end()

    # trailing text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = FONT_MAIN
        run.font.size = FONT_SIZE_BODY
        if base_bold:
            run.bold = True
        if base_italic:
            run.italic = True


# ──────────────────────────────────────────────────────────────────────────────
# Table builder
# ──────────────────────────────────────────────────────────────────────────────

def add_markdown_table(doc: Document, lines: list[str]):
    """Parse markdown table lines into a Word table."""
    # Filter out separator lines (|---|---|) — may contain |, -, :, spaces
    data_lines = [l for l in lines if not re.match(r"^[\s\-:|]+$", l.strip())]
    if not data_lines:
        return

    rows_data = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    n_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Auto-fit table to page width so columns are not clipped
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")  # 100% of page width
    tblPr.append(tblW)

    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            if ci >= n_cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            add_formatted_text(p, cell_text, doc)
            # Make header row bold + centered
            if ri == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                # Shading
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9E2F3")
                cell._tc.get_or_add_tcPr().append(shd)

    # Set font for all cells
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if not run.font.name:
                        run.font.name = FONT_MAIN
                    if not run.font.size:
                        run.font.size = Pt(10)

    doc.add_paragraph()  # spacer


# ──────────────────────────────────────────────────────────────────────────────
# Code block builder
# ──────────────────────────────────────────────────────────────────────────────

def add_code_block(doc: Document, code_lines: list[str], lang: str = ""):
    """Add a code block with gray background and monospace font."""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(line)
        run.font.name = FONT_CODE
        run.font.size = FONT_SIZE_CODE
        run.font.color.rgb = RGBColor(30, 30, 30)
        # gray background on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
    doc.add_paragraph()  # spacer


# ──────────────────────────────────────────────────────────────────────────────
# Main parser
# ──────────────────────────────────────────────────────────────────────────────

def process_file(doc: Document, filepath: Path):
    """Read a chapter .txt file and convert its Markdown-like content to docx."""
    text = filepath.read_text(encoding="utf-8-sig")  # utf-8-sig strips BOM if present
    lines = text.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Empty line → skip (inter-paragraph spacing handled by style) ──
        if not stripped:
            i += 1
            continue

        # ── Bibliography section → stop processing file (built separately) ──
        if stripped in ("## Literatura", "BIBLIOGRAFIA", "## Bibliografia"):
            break

        # ── [SCREEN: ...] placeholder → skip (images inserted via blockquotes) ──
        if stripped.startswith("[SCREEN:"):
            i += 1
            continue

        # ── Code block ``` ──
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            add_code_block(doc, code_lines, lang)
            continue

        # ── Headings ──
        if stripped.startswith("# "):
            global _first_h1_seen
            heading_text = stripped[2:].strip()
            # Remove markdown formatting markers from heading
            heading_text = re.sub(r"\*{1,3}", "", heading_text)
            heading_text = apply_nbsp(heading_text)
            p = doc.add_heading(heading_text, level=1)
            p.paragraph_format.first_line_indent = Cm(0)
            if not _first_h1_seen:
                p.paragraph_format.page_break_before = True
                _first_h1_seen = True
            else:
                p.paragraph_format.page_break_before = False
            i += 1
            continue

        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            heading_text = re.sub(r"\*{1,3}", "", heading_text)
            heading_text = apply_nbsp(heading_text)
            p = doc.add_heading(heading_text, level=2)
            p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        if stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            heading_text = re.sub(r"\*{1,3}", "", heading_text)
            heading_text = apply_nbsp(heading_text)
            p = doc.add_heading(heading_text, level=3)
            p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        # ── Display math (standalone $$ ... $$) ──
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            math_text = stripped[2:-2].strip()
            add_display_equation(doc, math_text)
            i += 1
            continue

        if stripped == "$$":
            # Multi-line display math
            math_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                math_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                i += 1
            add_display_equation(doc, " ".join(math_lines))
            continue

        # ── Markdown table ──
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        # ── Blockquote (figure/table captions) ──
        if stripped.startswith("> "):
            quote_parts: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_parts.append(lines[i].strip()[2:])
                i += 1
            full_quote = " ".join(quote_parts)

            # ── Insert image BEFORE caption (if this is a figure) ──
            fig_match = re.search(r"\*\*Rysunek\s+(\d+\.\d+)\*\*", full_quote)
            if fig_match:
                fig_num = fig_match.group(1)
                _insert_figure_image(doc, fig_num)
                # Collect caption for Spis ilustracji
                cap_text = re.sub(r"\*{1,3}", "", full_quote).strip()
                _collected_figures.append((fig_num, cap_text))

            # ── Collect table caption from blockquote ──
            tab_match = re.search(r"\*\*Tabela\s+(\d+\.\d+)", full_quote)
            if tab_match and not fig_match:
                tab_num = tab_match.group(1)
                cap_text = re.sub(r"\*{1,3}", "", full_quote).strip()
                _collected_tables.append((tab_num, cap_text))

            # ── Caption text (italic) ──
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_text(p, full_quote, doc, base_italic=True)

            # ── Source citation footnote for figures ──
            if fig_match:
                source = _read_figure_source(fig_num)
                if source:
                    site_name, url, access_date = source
                    footnote_text = (
                        f"Źródło ilustracji: {site_name}, {url}, "
                        f"dostęp: {access_date}."
                    )
                    add_footnote_ref(p, doc, footnote_text)
            continue

        # ── Numbered list (1. / 2. etc.) ──
        m_num = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.first_line_indent = Cm(0)
            add_formatted_text(p, m_num.group(2), doc)
            i += 1
            # Continuation lines (indented or starting with spaces)
            while i < len(lines):
                next_line = lines[i]
                next_stripped = next_line.strip()
                # continuation: indented line that is not a new item / heading / empty
                if next_stripped and not re.match(r"^\d+\.\s", next_stripped) \
                   and not next_stripped.startswith("#") \
                   and not next_stripped.startswith("|") \
                   and not next_stripped.startswith("```") \
                   and not next_stripped.startswith("> ") \
                   and not next_stripped.startswith("- ") \
                   and not next_stripped.startswith("* ") \
                   and (next_line.startswith("   ") or next_line.startswith("\t")):
                    add_formatted_text(p, " " + next_stripped, doc)
                    i += 1
                else:
                    break
            continue

        # ── Bullet list (- or *) ──
        m_bullet = re.match(r"^[-*]\s+(.+)", stripped)
        if m_bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            add_formatted_text(p, m_bullet.group(1), doc)
            i += 1
            # Continuation lines (indented)
            while i < len(lines):
                next_line = lines[i]
                next_stripped = next_line.strip()
                if next_stripped and not re.match(r"^[-*]\s", next_stripped) \
                   and not next_stripped.startswith("#") \
                   and not next_stripped.startswith("|") \
                   and not next_stripped.startswith("```") \
                   and not next_stripped.startswith("> ") \
                   and not re.match(r"^\d+\.\s", next_stripped) \
                   and (next_line.startswith("  ") or next_line.startswith("\t")):
                    add_formatted_text(p, " " + next_stripped, doc)
                    i += 1
                else:
                    break
            continue

        # ── Horizontal rule --- ──
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # ── Regular paragraph (may span multiple non-empty lines) ──
        para_parts: list[str] = []
        while i < len(lines):
            cl = lines[i]
            cs = cl.strip()
            # Stop at: empty line, actual heading (# + space), table, codeblock,
            # blockquote, list, hr, display math
            if not cs or re.match(r"^#{1,3}\s", cs) or cs.startswith("|") \
               or cs.startswith("```") or cs.startswith("> ") \
               or re.match(r"^[-*]\s", cs) or re.match(r"^\d+\.\s", cs) \
               or re.match(r"^-{3,}$", cs) or cs.startswith("$$"):
                break
            para_parts.append(cs)
            i += 1

        if para_parts:
            full_text = " ".join(para_parts)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_formatted_text(p, full_text, doc)

            # Collect table caption (bold line like **Tabela 4.1. ...**)
            tab_cap_match = re.search(r"\*\*Tabela\s+(\d+\.\d+)\.?\s+(.+?)\*\*", full_text)
            if tab_cap_match:
                tab_num = tab_cap_match.group(1)
                tab_title = tab_cap_match.group(2).strip()
                _collected_tables.append((tab_num, f"Tabela {tab_num}. {tab_title}"))
        else:
            # Safety: skip unrecognised line to avoid infinite loop
            i += 1


# ──────────────────────────────────────────────────────────────────────────────
# Entry point
# ──────────────────────────────────────────────────────────────────────────────

def main():
    # Ensure UTF-8 output on Windows consoles (avoids cp1252 encoding errors)
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    print("=== build_docx.py — Budowanie dokumentu Word ===\n")

    for f in INPUT_FILES:
        if not f.exists():
            print(f"  BŁĄD: Brak pliku {f.name}")
            sys.exit(1)
        print(f"  ✓ {f.name}  ({f.stat().st_size:,} B)")

    # Load bibliography so footnotes contain real citations
    _load_bibliography()
    print(f"  ✓ Załadowano {len(_bibliography)} pozycji bibliograficznych")

    doc = create_document(base_docx=SCRIPT_DIR / "TitlePage.docx")

    # Reset collectors
    _collected_figures.clear()
    _collected_tables.clear()

    # Process each chapter file
    for filepath in INPUT_FILES:
        print(f"\n  Przetwarzanie: {filepath.name} ...")
        process_file(doc, filepath)

    # Build back-matter sections
    _build_bibliography(doc)
    _build_list_of_figures(doc)
    _build_list_of_tables(doc)

    # Save
    doc.save(str(OUTPUT_FILE))
    size_kb = OUTPUT_FILE.stat().st_size / 1024
    print(f"\n  ✓ Zapisano: {OUTPUT_FILE.name}  ({size_kb:.0f} KB)")
    print(f"    Ścieżka: {OUTPUT_FILE}")
    print("\nGotowe!")


if __name__ == "__main__":
    main()
